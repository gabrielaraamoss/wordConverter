import os
from flask import Flask, send_file
from docx import Document
from docx.shared import Inches
from io import BytesIO
from bs4 import BeautifulSoup

app = Flask(__name__)

@app.route('/generar_docx')
def generar_docx():
    ruta_html = os.path.join(os.path.dirname(__file__), 'ejemplo.html')
    
    with open(ruta_html, 'r', encoding='utf-8') as file:
        html_content = file.read()

    soup = BeautifulSoup(html_content, 'html.parser')
    texto_html = soup.get_text()  

    doc = Document()

    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True  


    logo_left_path = os.path.join(os.path.dirname(__file__), 'assets', 'logo2.png')
    cell_left = table.cell(0, 0)
    cell_left.paragraphs[0].add_run().add_picture(logo_left_path, width=Inches(1)) 
    cell_left.paragraphs[0].alignment = 0  

    logo_right_path = os.path.join(os.path.dirname(__file__), 'assets', 'logo2.png')  
    cell_right = table.cell(0, 1)
    cell_right.paragraphs[0].add_run().add_picture(logo_right_path, width=Inches(1))  
    cell_right.paragraphs[0].alignment = 2  

        
    doc.add_paragraph(texto_html)

    for element in soup.find_all(['strong']):
        if element.name == 'strong':  
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(element.get_text()) 
            run.bold = True  

    docx_io = BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)

    docx_io.seek(0)
    temp_file_path = os.path.join(os.path.dirname(__file__), 'reporte_generado.docx')
    with open(temp_file_path, 'wb') as f:
        f.write(docx_io.read())

    return send_file(temp_file_path, as_attachment=True, download_name="reporte_generado.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == '__main__':
    app.run(debug=True)
